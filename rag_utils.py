import hashlib

import requests
from bs4 import BeautifulSoup
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi
from sentence_transformers import CrossEncoder


# Create embedding model
def create_embeddings():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )


# Create cross-encoder reranker
def create_reranker():
    return CrossEncoder(
        "cross-encoder/ms-marco-MiniLM-L-6-v2"
    )


# Extract text from PDF
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1
    ):
        page_text = page.extract_text()

        if page_text:
            pages.append({
                "text": page_text,
                "page": page_number
            })

    return pages


# Extract text from TXT
def extract_text_from_txt(txt_file):
    text = txt_file.getvalue().decode(
        "utf-8",
        errors="ignore"
    )

    if not text.strip():
        return []

    return [{
        "text": text,
        "page": 1
    }]


# Extract text from DOCX
def extract_text_from_docx(docx_file):
    document = DocxDocument(docx_file)

    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )

    if not text.strip():
        return []

    return [{
        "text": text,
        "page": 1
    }]


# Extract text from webpage
def extract_text_from_web(url):
    response = requests.get(
        url,
        timeout=15,
        headers={
            "User-Agent": "Mozilla/5.0"
        }
    )

    response.raise_for_status()

    soup = BeautifulSoup(
        response.text,
        "html.parser"
    )

    for element in soup(
        ["script", "style", "noscript"]
    ):
        element.decompose()

    text = soup.get_text(
        separator=" ",
        strip=True
    )

    if not text:
        return []

    return [{
        "text": text,
        "page": 1
    }]


# Split text into chunks
def split_text(pages):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )

    chunks = []

    for page in pages:

        page_chunks = text_splitter.split_text(
            page["text"]
        )

        for chunk in page_chunks:

            chunks.append({
                "text": chunk,
                "page": page["page"]
            })

    return chunks


# Create a unique collection ID
def create_collection_id(source_name):
    return hashlib.sha256(
        source_name.encode("utf-8")
    ).hexdigest()


# Create Chroma vector store
def create_vector_store(
    chunks,
    embedding_model,
    source_name,
    collection_id
):
    texts = [
        chunk["text"]
        for chunk in chunks
    ]

    metadatas = [
        {
            "page": chunk["page"],
            "source": source_name
        }
        for chunk in chunks
    ]

    return Chroma.from_texts(
        texts=texts,
        embedding=embedding_model,
        metadatas=metadatas,
        collection_name=f"research_{collection_id}"
    )


# Create BM25 documents for a source
def create_bm25_documents(chunks, source_name):
    documents = []

    for chunk in chunks:
        documents.append(
            Document(
                page_content=chunk["text"],
                metadata={
                    "page": chunk["page"],
                    "source": source_name
                }
            )
        )

    return documents


# Create BM25 index
def create_bm25_index(documents):
    tokenized_documents = [
        document.page_content.lower().split()
        for document in documents
    ]

    return BM25Okapi(tokenized_documents)


# Retrieve documents using BM25
def bm25_search(
    bm25_index,
    documents,
    query,
    top_k=5
):
    if not documents:
        return []

    tokenized_query = query.lower().split()

    scores = bm25_index.get_scores(
        tokenized_query
    )

    ranked_indices = sorted(
        range(len(scores)),
        key=lambda index: scores[index],
        reverse=True
    )

    return [
        documents[index]
        for index in ranked_indices[:top_k]
    ]


# Retrieve documents using hybrid retrieval
def search_multiple_documents(
    vector_stores,
    bm25_stores,
    query,
    top_k=3,
    candidate_k=5
):
    dense_results = []
    bm25_results = []

    # Dense vector retrieval
    for vector_store in vector_stores:

        results = (
            vector_store
            .similarity_search_with_score(
                query,
                k=candidate_k
            )
        )

        dense_results.extend(
            document
            for document, score in results
        )

    # BM25 keyword retrieval
    for bm25_store in bm25_stores:

        results = bm25_search(
            bm25_store["index"],
            bm25_store["documents"],
            query,
            top_k=candidate_k
        )

        bm25_results.extend(results)

    # Combine results and remove duplicates
    combined_documents = {}
    dense_rankings = {}
    bm25_rankings = {}

    for rank, document in enumerate(
        dense_results,
        start=1
    ):
        key = (
            document.metadata.get("source"),
            document.metadata.get("page"),
            document.page_content
        )

        combined_documents[key] = document
        dense_rankings[key] = rank

    for rank, document in enumerate(
        bm25_results,
        start=1
    ):
        key = (
            document.metadata.get("source"),
            document.metadata.get("page"),
            document.page_content
        )

        combined_documents[key] = document
        bm25_rankings[key] = rank

    # Reciprocal Rank Fusion
    hybrid_scores = {}

    for key in combined_documents:

        score = 0

        if key in dense_rankings:
            score += 1 / (60 + dense_rankings[key])

        if key in bm25_rankings:
            score += 1 / (60 + bm25_rankings[key])

        hybrid_scores[key] = score

    ranked_documents = sorted(
        combined_documents.values(),
        key=lambda document: hybrid_scores[
            (
                document.metadata.get("source"),
                document.metadata.get("page"),
                document.page_content
            )
        ],
        reverse=True
    )

    # Return candidates for reranking
    return ranked_documents[:candidate_k]


# Rerank retrieved documents using a cross-encoder
def rerank_documents(
    reranker,
    query,
    documents,
    top_k=3
):
    if not documents:
        return []

    pairs = [
        (query, document.page_content)
        for document in documents
    ]

    scores = reranker.predict(pairs)

    ranked_results = sorted(
        zip(documents, scores),
        key=lambda item: item[1],
        reverse=True
    )

    return [
        document
        for document, score in ranked_results[:top_k]
    ]